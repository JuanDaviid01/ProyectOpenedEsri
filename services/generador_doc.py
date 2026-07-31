import re
# pyrefly: ignore [missing-import]
import fitz
from pathlib import Path
# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Inches
# pyrefly: ignore [missing-import]
from docx.enum.table import WD_ALIGN_VERTICAL

def obtener_ruta_base():
    return Path(__file__).resolve().parent.parent

def cargar_plantilla(tipo_certificado="opened"):
    nombre_archivo = "resolucionlimpiaEsri.docx" if tipo_certificado.strip().lower() == "esri" else "resolucionlimpia.docx"
    ruta = obtener_ruta_base() / "documentos" / nombre_archivo
    if not ruta.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {ruta}")
    return Document(ruta)

def reemplazar_texto_en_parrafo(paragraph, reemplazos):
    if not paragraph.text.strip():
        return

    texto_original = paragraph.text
    texto_nuevo = texto_original

    for variable, valor in reemplazos.items():
        patron = r"\{\{\s*" + re.escape(variable) + r"\s*\}\}"
        texto_nuevo = re.sub(patron, str(valor), texto_nuevo)

    if texto_nuevo != texto_original:
        paragraph.text = texto_nuevo

def reemplazar_texto_en_documento(doc, reemplazos):
    for p in doc.paragraphs:
        reemplazar_texto_en_parrafo(p, reemplazos)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    reemplazar_texto_en_parrafo(p, reemplazos)

    for s in doc.sections:
        for h in [s.header, s.footer]:
            for p in h.paragraphs:
                reemplazar_texto_en_parrafo(p, reemplazos)
            for t in h.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            reemplazar_texto_en_parrafo(p, reemplazos)

def limpiar_filas_tabla(tabla, fila_inicial=2):
    while len(tabla.rows) > fila_inicial:
        tabla._tbl.remove(tabla.rows[-1]._tr)

def llenar_tabla_cursos(doc, cursos, tipo_certificado="opened"):
    if not doc.tables:
        return

    tabla = doc.tables[0]
    limpiar_filas_tabla(tabla, fila_inicial=2)
    es_esri = (tipo_certificado.strip().lower() == "esri") or (len(tabla.rows[0].cells) == 6)
    
    for i, curso in enumerate(cursos):
        fila = tabla.add_row().cells
        fila[1].text = str(curso.get("asignatura_opened", curso.get("curso", "")))
        fila[2].text = str(curso.get("nota", ""))
        
        if es_esri:
            if i == 0:
                fila[0].text = str(curso.get("periodo", ""))
                fila[3].text = str(curso.get("asignatura_homologada", ""))
                fila[4].text = str(curso.get("codigo_asignatura", ""))
                fila[5].text = str(curso.get("creditos", ""))
            else:
                fila[0].text = ""
                fila[3].text = ""
                fila[4].text = ""
                fila[5].text = ""
        else:
            if i == 0:
                fila[0].text = str(curso.get("periodo", ""))
                fila[3].text = str(curso.get("nota_definitiva", ""))
                fila[4].text = str(curso.get("asignatura_homologada", ""))
                fila[5].text = str(curso.get("codigo_asignatura", ""))
                fila[6].text = str(curso.get("creditos", ""))
            else:
                fila[0].text = ""
                fila[3].text = ""
                fila[4].text = ""
                fila[5].text = ""
                fila[6].text = ""

    num_rows = len(cursos)
    start_row = 2
    end_row = start_row + num_rows - 1

    if num_rows > 1:
        cols_to_merge = [0, 3, 4, 5] if es_esri else [0, 3, 4, 5, 6]
        for col_idx in cols_to_merge:
            base_cell = tabla.cell(start_row, col_idx)
            for r_idx in range(start_row + 1, end_row + 1):
                base_cell.merge(tabla.cell(r_idx, col_idx))
            base_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def anexar_pdfs(doc, rutas_pdfs_anexos, carpeta_salida):
    for ruta_pdf in rutas_pdfs_anexos:
        pdf_doc = fitz.open(ruta_pdf)
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            pix = page.get_pixmap(dpi=150)
            
            ruta_img_temp = Path(carpeta_salida) / f"temp_{Path(ruta_pdf).stem}_{page_num}.png"
            pix.save(str(ruta_img_temp))
            
            doc.add_page_break()
            doc.add_picture(str(ruta_img_temp), width=Inches(6.5))
            
        pdf_doc.close()

def generar_resolucion_opened(datos, cursos, carpeta_salida, nombre_salida, rutas_pdfs_anexos):
    return generar_resolucion_word(datos, cursos, carpeta_salida, nombre_salida, rutas_pdfs_anexos, tipo_certificado="opened")

def generar_resolucion_esri(datos, cursos, carpeta_salida, nombre_salida, rutas_pdfs_anexos):
    return generar_resolucion_word(datos, cursos, carpeta_salida, nombre_salida, rutas_pdfs_anexos, tipo_certificado="esri")

def generar_resolucion_word(datos, cursos, carpeta_salida, nombre_salida, rutas_pdfs_anexos, tipo_certificado="opened"):
    doc = cargar_plantilla(tipo_certificado=tipo_certificado)
    
    # 1. Inyección de payload de negocio
    reemplazar_texto_en_documento(doc, datos)
    llenar_tabla_cursos(doc, cursos, tipo_certificado=tipo_certificado)

    # 2. Rasterización de anexos PDF a PNG y anexado al DOM del DOCX
    anexar_pdfs(doc, rutas_pdfs_anexos, carpeta_salida)

    # 3. Guardado final
    ruta_docx = Path(carpeta_salida) / nombre_salida
    doc.save(ruta_docx)

    return str(ruta_docx)